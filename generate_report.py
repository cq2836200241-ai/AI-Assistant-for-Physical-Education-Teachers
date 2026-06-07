import os
import sys

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    print("python-docx is not installed.")
    sys.exit(1)

def create_report():
    doc = Document()

    # 标题
    title = doc.add_heading('智能体育教案助手 开发与应用报告', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 一、项目概述
    doc.add_heading('一、 项目概述', level=1)
    p = doc.add_paragraph()
    p.add_run('项目背景：').bold = True
    p.add_run('随着人工智能技术的普及，为了减轻体育教师日常备课的负担并提升教案编写的专业度，本项目应运而生。\n')
    p.add_run('项目定位：').bold = True
    p.add_run('一款基于大模型（AI驱动）的体育教案生成与管理辅助工具。\n')
    p.add_run('支持平台：').bold = True
    p.add_run('支持现代浏览器环境（Web端），并通过 Electron 封装支持跨平台桌面端应用。')

    # 二、技术架构与选型
    doc.add_heading('二、 技术架构与选型', level=1)
    p = doc.add_paragraph()
    p.add_run('项目采用了现代化的前端技术栈，具备优异的开发体验与运行性能：\n')
    
    bullet_points = [
        ('核心框架：', 'React 19 配合 Vite 6 构建，采用 TypeScript 提供类型安全。'),
        ('桌面端方案：', 'Electron 结合 electron-builder 进行跨平台应用打包与分发。'),
        ('状态管理：', '使用 Zustand 5 实现轻量且高效的全局状态管理。'),
        ('样式与UI：', 'Tailwind CSS v4 结合 Shadcn UI 构建美观、响应式的组件库。'),
        ('大模型引擎：', '接入 Google Gemini API（@google/genai），实现智能化的教案内容生成。'),
        ('后端服务与存储：', '集成 Firebase 用于数据存储与验证。'),
        ('文档导出模块：', '内置 docx、jspdf、html2pdf.js 等库，支持一键导出为 PDF 和 Word 格式。')
    ]
    
    for bold_text, normal_text in bullet_points:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(bold_text).bold = True
        p.add_run(normal_text)

    # 三、核心功能设计
    doc.add_heading('三、 核心功能设计', level=1)
    features = [
        ('1. 智能教案生成：', '系统能够根据用户填写的课题、年级、课型、能力等参数，动态构建 Prompt 并调用大模型，生成标准化、结构化的体育教案（包含准备部分、基本部分、结束部分等）。'),
        ('2. 丰富的教学资源库：', '内置“游戏库 (Game Library)”与“运动拆解 (Movement Analysis)”功能，为教师设计教学环节提供丰富的灵感与素材。'),
        ('3. 智能课表管理：', '提供“智能课表 (Smart Schedule)”功能，帮助教师直观地管理和规划日常教学进度。'),
        ('4. 多端自适应与沉浸式体验：', '针对移动端和桌面端提供不同的界面布局（如移动端抽屉式配置，桌面端左侧控制面板）。生成过程配有打字机动效和完成后纸屑动画（canvas-confetti）反馈。'),
        ('5. 教案管理与导出：', '支持教案的收藏（采纳的教案）与历史记录管理，一键预览并将其转化为专业排版的 Markdown 或直接下载为 Docx/PDF 文档。')
    ]
    for bold_text, normal_text in features:
        p = doc.add_paragraph()
        p.add_run(bold_text).bold = True
        p.add_run(normal_text)

    # 四、应用场景与业务价值
    doc.add_heading('四、 应用场景与业务价值', level=1)
    p = doc.add_paragraph()
    p.add_run('1. 极大提升备课效率：').bold = True
    p.add_run('将原本需要数小时撰写的教案缩短至几分钟的配置与微调，让教师将更多精力投入到实际教学互动中。\n')
    p.add_run('2. 规范化与标准化：').bold = True
    p.add_run('AI 生成的内容严格按照体育教学规范（结构、时间分配等）进行编排，保障教案质量。\n')
    p.add_run('3. 激发教学创新：').bold = True
    p.add_run('结合运动拆解与游戏库，帮助教师引入新的热身游戏与训练手段，增强课堂趣味性。')

    # 五、总结与展望
    doc.add_heading('五、 总结与展望', level=1)
    p = doc.add_paragraph('“智能体育教案助手”成功将大模型生成能力与垂直领域业务结合，通过完善的工程化实现与细腻的交互设计，打造了一款高可用、高颜值的应用。未来可以进一步引入个性化学生体质数据分析与更多专项运动模型，使其成为体育教育的全面数字化平台。')

    # 保存
    save_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '智能体育教案助手_开发与应用报告.docx')
    doc.save(save_path)
    print(f"Report successfully saved to: {save_path}")

if __name__ == '__main__':
    create_report()
